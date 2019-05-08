import pandas as pd
import os
import re
from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table,_Row
from docx.text.paragraph import Paragraph
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import Inches

class MSR:

    def __init__(self, tasknumber):
        self.taskOrder = tasknumber

    def readAllDocx(self,datafromexcel):
        path = r"C:\Users\vgollapudi\Desktop\MSR\all"
        outputDoc = Document()

        def iter_block_items(parent):
            if isinstance(parent, _Document):
                parent_elm = parent.element.body
            elif isinstance(parent, _Cell):
                parent_elm = parent._tc
            elif isinstance(parent, _Row):
                parent_elm = parent._tr
            else:
                raise ValueError("something's not right")
            for child in parent_elm.iterchildren():
                if isinstance(child, CT_P):
                    yield Paragraph(child, parent)
                elif isinstance(child, CT_Tbl):
                    yield Table(child, parent)

        def get_para_data(output_doc_name, paragraph):
            """
            Write the run to the new file and then set its font, bold, alignment, color etc. data.
            """

            output_para = output_doc_name.add_paragraph()
            for run in paragraph.runs:
                s=run.text
                s=re.sub('<[^>]+>', '', s)
                output_run = output_para.add_run(s)
                output_run.bold = run.bold
                output_run.italic = run.italic
                output_run.underline = run.underline
                output_run.font.color.rgb = run.font.color.rgb
                output_run.style.name = run.style.name
            output_para.paragraph_format.alignment = paragraph.paragraph_format.alignment
            if paragraph.style.name == "List Paragraph" and len(output_para.text) > 4:
                output_para.style = output_doc_name.styles['List Bullet']
            else:
                output_para.style = paragraph.style

        for fileName in datafromexcel.keys():
            filePath = os.path.join(path, fileName + ".docx")
            exists = os.path.isfile(filePath)
            if exists:
                input_doc = Document(filePath)
                counter = 0
                for block in iter_block_items(input_doc):
                    if isinstance(block, Paragraph):
                        if counter == 0:
                            output_para = outputDoc.add_paragraph()
                            output_para.add_run(fileName)
                            output_para.style = outputDoc.styles['Heading 2']
                            timesheetExcel = datafromexcel[fileName]
                            timesheetExcel = timesheetExcel.drop(timesheetExcel.columns[[4,5, 6, 8]], axis=1)
                            print(timesheetExcel.shape)
                            timesheetExcel= timesheetExcel.append(pd.DataFrame({"Labor Category":[fileName],"Current Hours":[sum(timesheetExcel["Current Hours"])]}),ignore_index = True,sort=False)
                            print(timesheetExcel.shape)
                            t = outputDoc.add_table(timesheetExcel.shape[0] + 1, timesheetExcel.shape[1])
                            t.style = 'Table Grid'
                            for j in range(timesheetExcel.shape[-1]):
                                shading_elm_1 = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls('w')))
                                t.cell(0, j).text = timesheetExcel.columns[j]
                                t.cell(0, j)._tc.get_or_add_tcPr().append(shading_elm_1)
                            for i in range(timesheetExcel.shape[0]):
                                for j in range(timesheetExcel.shape[-1]):
                                    t.cell(i + 1, j).text = str(timesheetExcel.values[i, j])
                                    if j == 0:
                                        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="dAEAF7"/>'.format(nsdecls('w')))
                                        t.cell(i + 1, j)._tc.get_or_add_tcPr().append(shading_elm_1)
                                        t.cell(i + 1, j).width = Inches(2)

                            counter += 1
                        else:
                            get_para_data(outputDoc, block)
                    elif isinstance(block, Table):
                        output_para = outputDoc.add_paragraph()
                        tab = block.table._tbl
                        output_para._p.addnext(tab)


        outputDoc.save(r'C:\Users\vgollapudi\Desktop\MSR\out\out.docx')

    def readExcel(self):
        path = r"C:\Users\vgollapudi\Desktop\MSR\TO1 and TO6.xlsx"
        # print(path)
        df = pd.read_excel(path, sheet_name='Invoice Detail '+self.taskOrder)
        individual_project = {}
        for x in range(0, df.shape[0]):
            if df.iloc[x, 0] == "Labor Category" and df.iloc[x, 1] == "Type" and df.iloc[x, 2] == "First Name" and \
                    df.iloc[x, 3] == "Last Name":
                final = df.loc[x + 1:, :]
                final.columns = list(df.iloc[x, :])
                break
        counter = 0
        #
        for x in range(0, final.shape[0]):
            if final.iloc[x, 3] == "Total":
                individual_project[final.iloc[x, 0]] = final.iloc[counter:x, :].dropna()
                counter = x + 1
        return individual_project

if __name__ == '__main__':
    msr = MSR("TO6")
    data_from_excel = msr.readExcel()
    print(data_from_excel.keys())
    msr.readAllDocx(data_from_excel)
