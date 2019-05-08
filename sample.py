# from docx import Document
# from docx.enum.text import WD_ALIGN_PARAGRAPH
# ndocument = Document()
# document = Document(r'C:\Users\vgollapudi\Desktop\MSR\input\first pages.docx')
# sec = document.sections[0]
# header = sec.header
#
# for paragraph in header.paragraphs:
#     print(paragraph.text)
# # print(sec.header)
# head=ndocument.sections[0].header
# foot=ndocument.sections[0].footer
#
# # h.is_linked_to_previous
# head.paragraphs[0].text='''Federal Aviation Administration (FAA)		Software Solution Delivery (SSD) Support Services
# Office of Information Services (AIT)		Contract:  DTFACT-17-D-00007
# '''
# foot.paragraphs[0].text='''Monthly status Report              Page|1                  August 2018'''
# foot.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
#
#
# p = ndocument.add_paragraph()
# r = p.add_run()
# r.add_text('Good Morning every body,This is my ')
# r.add_picture(r'C:\Users\vgollapudi\Desktop\MSR\faaLogo.png')
# r.add_text(' do you like it?')
# # h.add_paragraph()
# # h.paragraphs[1].text+=sec.header.paragraphs[1].text
# ndocument.save(r'C:\Users\vgollapudi\Desktop\MSR\demo.docx')

# import openpyxl
#
# path = r"C:\Users\vgollapudi\Desktop\MSR\TO1 and TO6.xlsx"
#
# wb_obj = openpyxl.load_workbook(path)
#
# sheet_obj = wb_obj.active
# max_col = sheet_obj.max_column
#
# # Loop will print all columns name
# for i in range(1, max_col + 1):
#     cell_obj = sheet_obj.cell(row=13, column=i)
#     print(cell_obj.value)


###############
# import os
# from docx import Document
# import re
# from docx.document import Document as _Document
# from docx.oxml.text.paragraph import CT_P
# from docx.oxml.table import CT_Tbl
#
# fileNames = ['TO6-SA Technical Support', 'TO6-SA Sprint 0', 'TO6-Reimbursable ToolSet', 'TO6-LAANC', 'TO6-Quasar NCA']
#
# path = r"C:\Users\vgollapudi\Desktop\MSR\all"
# outputDoc = Document()
# def get_para_data(output_doc_name, paragraph):
#     """
#     Write the run to the new file and then set its font, bold, alignment, color etc. data.
#     """
#     # if isinstance(parent, _Document):
#
#     output_para = output_doc_name.add_paragraph()
#     for run in paragraph.runs:
#         s=run.text
#         s=re.sub('<[^>]+>', '', s)
#         output_para.add_run(s)
#         print(s)
#         # Run's bold data
#         # output_run.bold = run.bold
#         # Run's italic data
#         # output_run.italic = run.italic
#         # Run's underline data
#         # output_run.underline = run.underline
#         # Run's color data
#         # output_run.font.color.rgb = run.font.color.rgb
#         # Run's font data
#         # output_run.style.name = run.style.name
#     # Paragraph's alignment data
#     output_para.paragraph_format.alignment = paragraph.paragraph_format.alignment
#     if paragraph.style.name == "List Paragraph":
#         output_para.style = output_doc_name.styles['List Bullet']
#     else:
#         output_para.style = paragraph.style
#     print(output_para.style.name)
#
# for fileName in fileNames:
#     filePath = os.path.join(path, fileName + ".docx")
#     # print(filePath)
#     exists = os.path.isfile(filePath)
#     if exists:
#         input_doc = Document(filePath)
#         # print(filePath)
#         for para in input_doc.paragraphs:
#             # print(para.style.name)
#             get_para_data(outputDoc, para)
#
# outputDoc.save(r'C:\Users\vgollapudi\Desktop\MSR\out\out.docx')
#

from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table,_Row
from docx.text.paragraph import Paragraph
def iter_block_items(parent):
    """
Generate a reference to each paragraph and table child within *parent*,
in document order. Each returned value is an instance of either Table or
Paragraph. *parent* would most commonly be a reference to a main
Document object, but also works for a _Cell object, which itself can
contain paragraphs and tables.
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    elif isinstance(parent, _Row):
        parent_elm = parent._tr
    else:
        raise ValueError("something's not right")
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)
document = Document(r"C:\Users\vgollapudi\Desktop\MSR\all\TO6-SA Technical Support.docx")
for block in iter_block_items(document):

    #print(block.text if isinstance(block, Paragraph) else '<table>')
    if isinstance(block, Paragraph):
        pass
        # print(block.text)
    elif isinstance(block, Table):
        for row in block.rows:
            row_data = []
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    row_data.append(paragraph.text)
            print("\t".join(row_data))